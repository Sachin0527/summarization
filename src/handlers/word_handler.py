from docx import Document
from .base_handler import BaseHandler


class WordHandler(BaseHandler):
    def summarize(self, input_file, output_file):
        try:
            self.logger.debug(f"Reading Word file: {input_file}")
            doc = Document(input_file)
            text = ""

            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            if not text.strip():
                self.logger.error(f"Failed to extract text from the Word document: {input_file}")
                raise ValueError(f"Failed to extract text from the Word document: {input_file}")

            summary = self.call_api(text)

            self.logger.debug(f"Writing summary in Word file: {output_file}")
            # Create a new Word document for the summary
            summary_doc = Document()
            summary_doc.add_paragraph(summary)
            summary_doc.save(output_file)

            self.logger.debug(f"Summary written successfully to {output_file}")
            return f"Summary written successfully to {output_file}"
        except Exception as e:
            msg = f"Error summarizing Word file: {input_file}, Error: {str(e)}"
            self.logger.error(msg)
            raise Exception(msg)