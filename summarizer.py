import google.generativeai as genai
from config import API_KEY, GENERATION_CONFIG
from docx import Document

# Configure the API key
genai.configure(api_key=API_KEY)

# Create the GenerativeModel instance
model = genai.GenerativeModel(
    model_name="gemini-1.5-flash",
    generation_config=GENERATION_CONFIG
)

# Function to summarize text
def summarize_text(input_text):
    chat_session = model.start_chat(history=[])
    response = chat_session.send_message(f"Summarize this: {input_text}")
    return response.text

# Function to handle large texts by splitting into chunks
def chunk_text(text, max_tokens=1000):
    paragraphs = text.split("\n\n")  # Split based on paragraph breaks
    chunks = []
    current_chunk = ""

    for paragraph in paragraphs:
        if len(current_chunk) + len(paragraph) < max_tokens:
            current_chunk += paragraph + "\n\n"
        else:
            chunks.append(current_chunk.strip())
            current_chunk = paragraph + "\n\n"

    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks

# Function to summarize large text
def summarize_large_text(input_text):
    chunks = chunk_text(input_text)
    full_summary = []

    for idx, chunk in enumerate(chunks):
        print(f"Summarizing chunk {idx + 1} of {len(chunks)}...")
        summary = summarize_text(chunk)
        full_summary.append(summary)

    return "\n\n".join(full_summary)

# Function to read from a DOCX file
def read_from_docx(input_file_path):
    doc = Document(input_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n\n".join(full_text)

# Function to write summarized text to a DOCX file
def write_to_docx(output_file_path, summarized_text):
    doc = Document()
    doc.add_paragraph(summarized_text)
    doc.save(output_file_path)

# Function to summarize from a Word document and output to another Word document
def summarize_from_word(input_file_path, output_file_path):
    input_text = read_from_docx(input_file_path)
    summary = summarize_large_text(input_text)
    write_to_docx(output_file_path, summary)
    print(f"Summary has been written to {output_file_path}")

# Function to summarize text from a plain text file and write the summary to an output file
def summarize_from_file(input_file_path, output_file_path):
    with open(input_file_path, 'r', encoding='utf-8') as file:
        input_text = file.read().strip()

    summary = summarize_large_text(input_text)

    with open(output_file_path, 'w', encoding='utf-8') as file:
        file.write(summary)

# Determine file type and summarize accordingly
def summarize(input_file_path, output_file_path):
    if input_file_path.endswith('.docx'):
        summarize_from_word(input_file_path, output_file_path)
    else:
        summarize_from_file(input_file_path, output_file_path)

# Specify the paths to the input and output files
input_file_path = 'input.docx'  # or 'input.txt'
output_file_path = 'summary.docx'  # or 'summary.txt'

# Perform the summarization
summarize(input_file_path, output_file_path)