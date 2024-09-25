from docx import Document
from .summarization import summarize_large_text

def read_from_docx(input_file_path):
    doc = Document(input_file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n\n".join(full_text)

def write_to_docx(output_file_path, summarized_text):
    doc = Document()
    doc.add_paragraph(summarized_text)
    doc.save(output_file_path)

def summarize_docx(input_file_path, output_file_path):
    input_text = read_from_docx(input_file_path)
    summary = summarize_large_text(input_text)
    write_to_docx(output_file_path, summary)
    print(f"Summary has been written to {output_file_path}")
