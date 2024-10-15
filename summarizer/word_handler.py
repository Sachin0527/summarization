from docx import Document
from summarizer.summarization import summarize_large_text
from summarizer.logger import logger  # Import logger

def summarize_word(input_file_path, output_file_path):
    logger.info(f"Reading DOCX file: {input_file_path}")
    
    try:
        doc = Document(input_file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"

        if not text.strip():
            logger.error(f"Failed to extract text from the Word document: {input_file_path}")
            raise ValueError("No text found in the Word document.")
        
        logger.info(f"Word document text extracted successfully from {input_file_path}. Starting summarization.")
        summary = summarize_large_text(text)

        # Create a new Word document for the summary
        summary_doc = Document()
        summary_doc.add_paragraph(summary)
        summary_doc.save(output_file_path)

        logger.info(f"Summary written successfully to {output_file_path}")
    except Exception as e:
        logger.error(f"Error summarizing Word document: {input_file_path}, Error: {str(e)}")
        raise
